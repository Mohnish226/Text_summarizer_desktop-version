#import nltk
#import file_reader
#import docx2txt
from docx import Document

#document = Document('/Users/Mohnish_Devadiga/Desktop/Project-Status-Report.doc')
document = Document('/Users/Mohnish_Devadiga/Desktop/test.docx')
for para in document.paragraphs:
    print(para.text)
#print(nltk.corpus.gutenberg.fileids())

'''
for names in nltk.corpus.gutenberg.fileids():
    print(names)

print(pdfread.read_pdf("/Users/Mohnish_Devadiga/Desktop/Data/resume.pdf"))
print(pdfread.read_pdf("/Users/Mohnish_Devadiga/Desktop/Mohnish/Mohnish-Devadiga-online.pdf"))
'''
'''
file_reader.read_file("/Users/Mohnish_Devadiga/Desktop/Data/resume.pdf")
file_reader.read_docx("/Users/Mohnish_Devadiga/Desktop/Project-Status-Report.doc")
'''
#text = docx2txt.process("/Users/Mohnish_Devadiga/Desktop/Project-Status-Report.doc")